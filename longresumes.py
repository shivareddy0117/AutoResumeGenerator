from openai import OpenAI
from dotenv import load_dotenv
import os
import docx

# Load environment variables
load_dotenv()

# Get API Key from environment variables
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# Initialize the OpenAI client with your API key
client = OpenAI(api_key=OPENAI_API_KEY)

# Function to read a DOCX file and return its content
##################################################################################################
from docx import Document
import os
import re


def extract_resume_summary(filepath):
    doc = Document(filepath)
    text = '\n'.join([para.text for para in doc.paragraphs])

    # Define regex patterns for sections
    patterns = {
        "professional_summary": r"PROFESSIONAL SUMMARY:\s*((?:.|\n)*?)(?=\n[A-Z]|$)",

        
    }

    # Dictionary to hold the content of each section
    sections = {}

    # Extract each section using regex
    for key, pattern in patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        sections[key] = match.group(1).strip() if match else "Section not found"

    return sections

# Example usage
file_path = './careersoft_MR.docx'
resume_summary = extract_resume_summary(file_path)

# Print each section content stored in variables
# print("Professional Summary:")
summary_text = resume_summary['professional_summary']


def extract_resume_sections(file_path):
    document = Document(file_path)
    
    sections = {
        'Summary_Text': '',
        'Technical_Skills': '',
        'Capital_One_Experience': '',
        'Horizon_Experience': '',
        'OrangeScape_Experience': '',
        'Rational_Technologies_Experience': '',
        'Education_Details': ''
    }
    
    current_section = None
    experience_section = None
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        
        if 'SUMMARY' in text.upper():
            current_section = 'Summary_Text'
            continue  # Skip the header line
        elif 'TECHNICAL SKILLS' in text.upper():
            current_section = 'Technical_Skills'
            continue  # Skip the header line
        elif 'EXPERIENCE' in text.upper():
            current_section = 'Experience'
            continue  # Skip the header line
        elif 'EDUCATION' in text.upper():
            current_section = 'Education_Details'
            continue  # Skip the header line
        
        if current_section == 'Experience':
            if 'Capital One' in text:
                experience_section = 'Capital_One_Experience'
                continue  # Skip the company name line
            elif 'Horizon BCBS' in text:
                experience_section = 'Horizon_Experience'
                continue  # Skip the company name line
            elif 'OrangeScape' in text:
                experience_section = 'OrangeScape_Experience'
                continue  # Skip the company name line
            elif 'Rational Technologies' in text:
                experience_section = 'Rational_Technologies_Experience'
                continue  # Skip the company name line
        
        if current_section == 'Experience' and experience_section:
            sections[experience_section] += text + '\n'
        elif current_section in sections:
            sections[current_section] += text + '\n'
    
    # Trim any leading/trailing whitespace
    for key in sections:
        sections[key] = sections[key].strip()
    
    return sections

# Usage
# Replace this with the actual path to your resume file
file_path = r'C:\Users\SAHITHYAMOGILI\Desktop\Projects\Auto_Resume_Generator\careersoft_MR.docx'

# Check if the file exists
if not os.path.exists(file_path):
    print(f"Error: The file '{file_path}' does not exist.")
else:
  
    resume_sections = extract_resume_sections(file_path)


    # # To access technical skills
    technical_skills_text = resume_sections['Technical_Skills'] 
    experience1_text = resume_sections['Capital_One_Experience']
    experience2_text = resume_sections['Horizon_Experience']
    experience3_text = resume_sections['OrangeScape_Experience']
    experience4_text = resume_sections['Rational_Technologies_Experience']
    educational_details_text = resume_sections['Education_Details']


##################################################################################################

# Define a job description
job_descriptions = ["""Remote - Sr Data Scientist or AI/ML Engineer
Role: Sr Data Scientist or AI/ML Engineer

REMOTE

Long term contract

 

Skill: SAS, Python, Pyspark, Machine Learning and Strong/ recent Healthcare experience
RAG, LLAMA2/3, AI models, LLMs using tools such as Gradio and streamlit
Demonstrated understanding of analytical processes and approaches for at least one class of applications.
8+ years in data management and analysis using Python / Pyspark
Strong proficiency with data manipulation in python
Strong understanding of normalized/dimensional data modeling principles.
Strong Knowledge in using and developing in SAS.
2+ years of experience working as a Machine Learning Engineer and has implemented and deployed at least 2 large projects related to ML & Natural language processing (NLP)
Broad knowledge of existing data mining, NLP algorithms and technologies (ML & DL), working knowledge of related frameworks and technologies
Work closely with other development team members to understand complex product requirements and translate them into data engineering and/or data management designs.
Hands on experience programming languages in Python, Spark, SQL, Data frames and PySpark
5+ years of experience working within cloud ecosystem (Azure preferred, AWS, or GCP)


Thank you, 
Anvesh
469-887-4159
anvesh@quiddityinfotech.com

Quiddity Infotech LLC"""]

# Function to update resume based on job description
def updated_text(summary_text, experience1_text, technical_skills_text,  Job_description):
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": f"Based on the given  job description {job_description} I would like to upadate the summary of the resume {summary_text}, capital_one_experience {experience1_text}, and technical skills {technical_skills_text} respectively"},
            {"role": "user", "content": f"please give me the update text make sure that the experiences are quantified. also provide the updated summary, technical skills, captial_one_experience "},
            # {"role": "system", "content": "professor "},
            # {"role": "user", "content": job_description_text},
            # {"role": "system", "content": "Please rewrite the resume emphasizing the skills and experiences that are most relevant to the job description, and quantize the achievements."}
        
        ]
    )
    if response.choices:
        return response.choices[0].message.content
    return "No response generated."


for idx, job_description in enumerate(job_descriptions):
    updated_resume_text = updated_text(summary_text,experience1_text, technical_skills_text, job_description)
    #collection.insert_one({"job_description": job_description, "updated_resume": updated_resume})

    # Write to a text file
    with open(f"./job_{idx+1}.txt", "w") as file:
        file.write(f"job_no: {idx+1}\n\nupdated:\n{updated_resume_text}")





